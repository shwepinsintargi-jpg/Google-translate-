import streamlit as st
from googletrans import Translator
import PyPDF2
from docx import Document
from io import BytesIO
import time

# Page UI Setup
st.set_page_config(page_title="Resume-able Translator", layout="centered")
st.title("🌐 Smart PDF Translator")
st.info("Internet ပြတ်တောက်သွားပါကလည်း ရပ်တန့်သွားသော စာမျက်နှာမှ ပြန်စနိုင်ပါသည်။")

# Session State များ ကြေညာခြင်း (Resume လုပ်ရန်အတွက်)
if 'current_page' not in st.session_state:
    st.session_state.current_page = 0
if 'translated_texts' not in st.session_state:
    st.session_state.translated_texts = []
if 'is_processing' not in st.session_state:
    st.session_state.is_processing = False

translator = Translator()

uploaded_file = st.file_uploader("PDF ဖိုင် တင်ပါ", type="pdf")

if uploaded_file:
    pdf_reader = PyPDF2.PdfReader(uploaded_file)
    total_pages = len(pdf_reader.pages)
    
    # Progress Bar နှင့် Status
    prog_bar = st.progress(st.session_state.current_page / total_pages if total_pages > 0 else 0)
    status_msg = st.empty()

    col1, col2 = st.columns(2)
    
    with col1:
        if st.button("ဘာသာပြန်ခြင်း စတင်/ဆက်လုပ်မည်"):
            st.session_state.is_processing = True

    with col2:
        if st.button("ခေတ္တရပ်နားမည်"):
            st.session_state.is_processing = False
            st.warning(f"စာမျက်နှာ {st.session_state.current_page} တွင် ရပ်နားထားသည်။")

    # ဘာသာပြန်ခြင်း လုပ်ငန်းစဉ်
    if st.session_state.is_processing and st.session_state.current_page < total_pages:
        for i in range(st.session_state.current_page, total_pages):
            if not st.session_state.is_processing:
                break
                
            page = pdf_reader.pages[i]
            text = page.extract_text()
            
            if text:
                try:
                    # Quality အတွက် တစ်ကြောင်းချင်းပြန်ခြင်းနှင့် Delay ထည့်ခြင်း
                    lines = text.split('\n')
                    translated_lines = []
                    for line in lines:
                        if line.strip():
                            res = translator.translate(line, src='en', dest='my')
                            translated_lines.append(res.text)
                            time.sleep(0.3) # API Safety Delay
                    
                    page_result = "\n".join(translated_lines)
                    st.session_state.translated_texts.append((f"Page {i+1}", page_result))
                    
                    # နောက်တစ်မျက်နှာသို့ ကူးရန် မှတ်သားခြင်း
                    st.session_state.current_page = i + 1
                    prog_bar.progress(st.session_state.current_page / total_pages)
                    status_msg.success(f"✅ စာမျက်နှာ {i+1} ပြီးစီးပါပြီ")
                    
                except Exception as e:
                    st.session_state.is_processing = False
                    st.error(f"အင်တာနက်ပြတ်တောက်သွားပါသည် သို့မဟုတ် Error တက်သွားပါသည်။ ကျေးဇူးပြု၍ ခဏနေပြန်ကြိုးစားပါ။")
                    break

    # Download Button (ဘာသာပြန်ပြီးသမျှ စာမျက်နှာများကို Word ပြောင်းရန်)
    if st.session_state.translated_texts:
        doc = Document()
        for title, content in st.session_state.translated_texts:
            doc.add_heading(title, level=2)
            doc.add_paragraph(content)
        
        bio = BytesIO()
        doc.save(bio)
        st.markdown("---")
        st.download_button("📥 ဘာသာပြန်ပြီးသမျှဖိုင်ကို Word ဖြင့်ရယူရန်", 
                           data=bio.getvalue(), 
                           file_name="Translated_Progress.docx")
